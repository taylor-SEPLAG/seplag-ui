from pathlib import Path
import sys

from docx import Document


def main() -> None:
    source = Path(sys.argv[1])
    document = Document(source)

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            print(f"P{index:03d} [{paragraph.style.name}] {text}")

    for table_index, table in enumerate(document.tables, start=1):
        print(f"TABLE {table_index}")
        for row_index, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            print(f"R{row_index:03d} | " + " | ".join(cells))

    for section_index, section in enumerate(document.sections, start=1):
        header = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if header:
            print(f"HEADER {section_index}: {header}")
        if footer:
            print(f"FOOTER {section_index}: {footer}")


if __name__ == "__main__":
    main()
